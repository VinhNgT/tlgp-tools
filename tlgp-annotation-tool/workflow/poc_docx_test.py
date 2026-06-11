"""
poc_docx_test.py

Proof-of-concept: Generate a sample .docx with all TLGP spec doc formatting
to validate copy-paste fidelity into Google Docs.

Usage:
    uv run --with python-docx python poc_docx_test.py

Output: poc_sample.docx in the same directory.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ============================================================
# Style Constants (from workflow Appendix A)
# ============================================================

FONT_FAMILY = "Times New Roman"
FONT_SIZE_DEFAULT = Pt(12)
FONT_SIZE_API = Pt(10.5)

HEADING_COLOR = RGBColor(0x4F, 0x81, 0xBD)
BORDER_COLOR_HEX = "7F7F7F"
HEADER_BG_HEX = "F2F2F2"

TABLE_TOTAL_WIDTH_PT = 468  # 6.5 inches

INFO_COLS_PT = [120, 348]
UI_COLS_PT = [35, 100, 57, 50, 50, 50, 126]
INTERACTION_COLS_PT = [164.5, 303.5]
API_COLS_PT = [110, 108, 45, 83, 62, 60]

CELL_PAD_TOP_PT = 5
CELL_PAD_BOTTOM_PT = 5
CELL_PAD_LEFT_PT = 6
CELL_PAD_RIGHT_PT = 6

BORDER_WIDTH_PT = 0.75


# ============================================================
# Helpers
# ============================================================

def pt_to_emu(pt: float) -> int:
    """Convert points to EMU (English Metric Units)."""
    return int(pt * 12700)


def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set cell borders using XML manipulation."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")} />')
        tcPr.append(tcBorders)

    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val is None:
            continue
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} '
            f'w:val="single" '
            f'w:sz="{int(val["width"] * 8)}" '  # sz is in 1/8 pt
            f'w:color="{val["color"]}" '
            f'w:space="0"/>'
        )
        existing = tcBorders.find(qn(f'w:{edge}'))
        if existing is not None:
            tcBorders.remove(existing)
        tcBorders.append(element)


def set_cell_padding(cell, top_pt=0, bottom_pt=0, left_pt=0, right_pt=0):
    """Set cell padding (margins) using XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")} />')
        tcPr.append(tcMar)
    else:
        tcMar.clear()

    for edge, pt_val in [('top', top_pt), ('bottom', bottom_pt),
                          ('start', left_pt), ('end', right_pt)]:
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:w="{int(pt_val * 20)}" w:type="dxa"/>'
        )
        tcMar.append(el)


def set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    existing = tcPr.find(qn('w:shd'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shading)


def set_column_width(column_cells, width_pt: float):
    """Set width for all cells in a column."""
    width_emu = pt_to_emu(width_pt)
    for cell in column_cells:
        cell.width = width_emu


def set_cell_width(cell, width_pt: float):
    """Set width for a single cell."""
    cell.width = pt_to_emu(width_pt)


def apply_default_borders(cell):
    """Apply the standard border to a cell."""
    border_spec = {"width": BORDER_WIDTH_PT, "color": BORDER_COLOR_HEX}
    set_cell_border(cell, top=border_spec, bottom=border_spec,
                    left=border_spec, right=border_spec)


def apply_default_padding(cell):
    """Apply the standard cell padding."""
    set_cell_padding(cell, CELL_PAD_TOP_PT, CELL_PAD_BOTTOM_PT,
                     CELL_PAD_LEFT_PT, CELL_PAD_RIGHT_PT)


def set_cell_paragraph_spacing(cell, space_above_pt=3, space_below_pt=3):
    """Set paragraph spacing within a cell."""
    for para in cell.paragraphs:
        pf = para.paragraph_format
        pf.space_before = Pt(space_above_pt)
        pf.space_after = Pt(space_below_pt)


def style_cell_text(cell, text: str, bold=False, font_size=None,
                    alignment=None, font_name=None):
    """Set text and styling for a cell."""
    cell.text = ""
    para = cell.paragraphs[0]
    if alignment:
        para.alignment = alignment
    run = para.add_run(text)
    run.font.name = font_name or FONT_FAMILY
    run.font.size = font_size or FONT_SIZE_DEFAULT
    if bold:
        run.font.bold = True


def style_table(table, col_widths_pt: list, font_size=None):
    """Apply full styling to a table: borders, padding, column widths, header row."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths and apply borders/padding to all cells
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            set_cell_width(cell, col_widths_pt[col_idx])
            apply_default_borders(cell)
            apply_default_padding(cell)
            set_cell_paragraph_spacing(cell)

            # Header row styling
            if row_idx == 0:
                set_cell_shading(cell, HEADER_BG_HEX)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.bold = True
                        run.font.name = FONT_FAMILY
                        run.font.size = font_size or FONT_SIZE_DEFAULT

            # Set font for data rows
            if row_idx > 0:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = FONT_FAMILY
                        run.font.size = font_size or FONT_SIZE_DEFAULT


# ============================================================
# Table Builders
# ============================================================

def build_info_table(doc, label: str, description: str):
    """Build a 2×2 Info Table."""
    data = [
        ["Tên chức năng", f"Component {label}"],
        ["Mô tả", description],
    ]
    table = doc.add_table(rows=2, cols=2)
    for r, row_data in enumerate(data):
        for c, text in enumerate(row_data):
            style_cell_text(table.cell(r, c), text,
                            bold=(r == 0),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER if r == 0 else None)
    style_table(table, INFO_COLS_PT)
    return table


def build_screen_info_table(doc):
    """Build an N×2 Screen General Info Table."""
    data = [
        ["Mô tả chức năng", "Màn hình hiển thị thông tin chi tiết sản phẩm"],
        ["Người thực hiện (Actor)", "Người dùng"],
        ["Điều kiện đầu vào (Precondition)", "Đã đăng nhập\nĐã chọn sản phẩm từ danh sách"],
        ["Điều kiện kích hoạt chức năng (Trigger)", "Click vào sản phẩm từ danh sách"],
        ["Luồng thông thường (Main flow)", "Hiển thị thông tin chi tiết sản phẩm"],
        ["Điều kiện đầu ra (Post condition)", "Hiển thị đầy đủ thông tin sản phẩm"],
        ["Business Rule", "Không có"],
    ]
    table = doc.add_table(rows=len(data), cols=2)
    for r, row_data in enumerate(data):
        for c, text in enumerate(row_data):
            style_cell_text(table.cell(r, c), text,
                            bold=(r == 0),
                            alignment=WD_ALIGN_PARAGRAPH.CENTER if r == 0 else None)
    style_table(table, INFO_COLS_PT)
    return table


def build_ui_elements_table(doc):
    """Build a 7-column UI Elements Table."""
    headers = ["STT", "Tên", "Loại Control", "Bắt buộc", "Độ dài tối đa",
               "Chỉnh sửa", "Mô tả"]
    data = [
        ["1", "Nút quay lại", "Icon", "", "", "", "Quay về màn hình trước"],
        ["2", "Tiêu đề màn hình", "Text", "", "", "", "Hiển thị tên màn hình"],
        ["3", "Nút chia sẻ", "Icon", "", "", "", "Chia sẻ sản phẩm"],
        ["4", "Nút giỏ hàng", "Icon", "", "", "", "Mở giỏ hàng"],
    ]
    table = doc.add_table(rows=1 + len(data), cols=7)

    # Header row
    for c, text in enumerate(headers):
        style_cell_text(table.cell(0, c), text, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for r, row_data in enumerate(data):
        for c, text in enumerate(row_data):
            style_cell_text(table.cell(r + 1, c), text)

    style_table(table, UI_COLS_PT)
    return table


def build_interaction_table(doc):
    """Build a 2-column Interaction Events Table."""
    headers = ["Hành động của tác nhân", "Phản ứng của hệ thống"]
    data = [
        ["Click vào nút Back", "Quay về màn trước"],
        ["Click vào nút Chia sẻ", "Mở bottom sheet chia sẻ sản phẩm"],
        ["Click vào nút Giỏ hàng", "Chuyển đến màn hình giỏ hàng"],
    ]
    table = doc.add_table(rows=1 + len(data), cols=2)

    for c, text in enumerate(headers):
        style_cell_text(table.cell(0, c), text, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for r, row_data in enumerate(data):
        for c, text in enumerate(row_data):
            style_cell_text(table.cell(r + 1, c), text)

    style_table(table, INTERACTION_COLS_PT)
    return table


def build_api_table(doc, is_request=True):
    """Build a 6-column API Table."""
    headers = ["Tên tham số", "Ý nghĩa", "Bắt buộc", "Kiểu dữ liệu",
               "Giới hạn", "Giá trị mặc định"]

    if is_request:
        data = [
            ["merchant_id", "Mã đối tác", "Có", "String", "", ""],
            ["product_id", "Mã sản phẩm", "Có", "String", "", ""],
        ]
    else:
        data = [
            ["product_id", "Mã sản phẩm", "", "String", "", ""],
            ["product_name", "Tên sản phẩm", "", "String?", "", ""],
            ["price", "Giá sản phẩm", "", "double", "", ""],
            ["items[].sku_id", "Mã SKU", "", "String", "", ""],
            ["items[].stock", "Số lượng tồn kho", "", "int", "", ""],
            ["page_info.has_next_page", "Còn trang tiếp", "", "bool", "", ""],
        ]

    table = doc.add_table(rows=1 + len(data), cols=6)

    for c, text in enumerate(headers):
        style_cell_text(table.cell(0, c), text, bold=True,
                        font_size=FONT_SIZE_API,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER)

    for r, row_data in enumerate(data):
        for c, text in enumerate(row_data):
            style_cell_text(table.cell(r + 1, c), text, font_size=FONT_SIZE_API)

    style_table(table, API_COLS_PT, font_size=FONT_SIZE_API)
    return table


# ============================================================
# Heading Builders
# ============================================================

def add_h3_heading(doc, text: str):
    """Add a HEADING_3 paragraph: bold, #4F81BD."""
    para = doc.add_heading(level=3)
    run = para.add_run(text)
    run.font.name = FONT_FAMILY
    run.font.size = FONT_SIZE_DEFAULT
    run.font.color.rgb = HEADING_COLOR
    run.font.bold = True
    # Ensure East Asian font is also set
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_FAMILY)
    return para


def add_h4_heading(doc, text: str):
    """Add a HEADING_4 paragraph: italic, #4F81BD."""
    para = doc.add_heading(level=4)
    run = para.add_run(text)
    run.font.name = FONT_FAMILY
    run.font.size = FONT_SIZE_DEFAULT
    run.font.color.rgb = HEADING_COLOR
    run.font.italic = True
    run.font.bold = False
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_FAMILY)
    return para


def add_bold_normal_text(doc, text: str):
    """Add a bold normal text paragraph (for API headings)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = FONT_FAMILY
    run.font.size = FONT_SIZE_DEFAULT
    run.font.bold = True
    return para


def add_normal_text(doc, text: str):
    """Add a plain normal text paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = FONT_FAMILY
    run.font.size = FONT_SIZE_DEFAULT
    return para


# ============================================================
# Main: Build PoC Document
# ============================================================

def create_poc_document():
    doc = Document()

    # Set default font for the document
    style = doc.styles['Normal']
    style.font.name = FONT_FAMILY
    style.font.size = FONT_SIZE_DEFAULT

    # Set page margins to match standard (1 inch = 72pt all around)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ---- Component Section ----
    add_h3_heading(doc, "1.1.1. Component tiêu đề màn hình chi tiết sản phẩm")

    add_h4_heading(doc, "1.1.1.1 Thông tin chung về chức năng")
    build_info_table(doc, "tiêu đề màn hình chi tiết sản phẩm",
                     "Mô tả tiêu đề màn hình chi tiết sản phẩm")

    add_h4_heading(doc, "1.1.1.2 Màn hình chức năng")
    # Placeholder text for where image would go
    add_normal_text(doc, "[Annotated component image would be inserted here]")

    add_h4_heading(doc, "1.1.1.3 Mô tả chi tiết các thành phần trên màn hình")
    build_ui_elements_table(doc)

    add_h4_heading(doc, "1.1.1.4 Xử lý luồng sự kiện tương tác (3 bước)")
    build_interaction_table(doc)

    # ---- Screen Overview Section ----
    add_h3_heading(doc, "1.1.2. Màn hình Chi tiết sản phẩm")

    add_h4_heading(doc, "1.1.2.1 Thông tin chung về chức năng")
    build_screen_info_table(doc)

    add_h4_heading(doc, "1.1.2.2 Màn hình chức năng")
    add_normal_text(doc, "[Root annotated screenshot would be inserted here]")

    add_h4_heading(doc, "1.1.2.3 Mô tả chi tiết các thành phần trên màn hình")
    build_ui_elements_table(doc)

    add_h4_heading(doc, "1.1.2.4 Xử lý luồng sự kiện tương tác (3 bước)")
    build_interaction_table(doc)

    # ---- API Documentation Section ----
    add_bold_normal_text(doc, "1. GET Chi tiết sản phẩm")
    add_normal_text(doc, "URL: /mall-catalog/api/v1/products/{productID}")

    add_bold_normal_text(doc, "Request")
    build_api_table(doc, is_request=True)

    add_bold_normal_text(doc, "Response (data = ProductDetailDTO)")
    build_api_table(doc, is_request=False)

    add_normal_text(doc, "PriceTiersDTO fields (price_tiers)")
    build_api_table(doc, is_request=False)

    return doc


def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(script_dir, "poc_sample.docx")

    print("Generating PoC .docx...")
    doc = create_poc_document()
    doc.save(output_path)
    print(f"✅ Saved to: {output_path}")
    print()
    print("Next steps:")
    print("  1. Open poc_sample.docx locally to verify formatting")
    print("  2. Upload to Google Drive (drag-and-drop)")
    print("  3. Open the converted Google Doc")
    print("  4. Select All → Copy → Paste into an existing Google Doc")
    print("  5. Check: heading colors, table widths, borders, header bg, fonts")


if __name__ == "__main__":
    main()
