"""Build styled tables for each TLGP table type.

Uses python-docx with XML manipulation for properties not exposed
by the high-level API (borders, cell padding, shading).
"""

from __future__ import annotations

from docx.document import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt
from docx.table import Table

from tlgp_doc_generator.models import (
    ApiParam,
    ChildElement,
    Interaction,
)
from tlgp_doc_generator.style_constants import (
    API_COLS_PT,
    API_TABLE_HEADERS,
    BORDER_COLOR_HEX,
    BORDER_WIDTH_PT,
    CELL_PAD_BOTTOM_PT,
    CELL_PAD_LEFT_PT,
    CELL_PAD_RIGHT_PT,
    CELL_PAD_TOP_PT,
    CELL_SPACE_ABOVE_PT,
    CELL_SPACE_BELOW_PT,
    FONT_FAMILY,
    FONT_SIZE_API,
    FONT_SIZE_DEFAULT,
    HEADER_BG_HEX,
    INFO_COLS_PT,
    INTERACTION_COLS_PT,
    INTERACTION_TABLE_HEADERS,
    UI_COLS_PT,
    UI_TABLE_HEADERS,
)


# ============================================================
# Low-level XML helpers
# ============================================================


def _pt_to_emu(pt: float) -> int:
    """Convert points to EMU (English Metric Units)."""
    return int(pt * 12700)


def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set cell borders via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")} />')
        tcPr.append(tcBorders)

    for edge, val in [
        ("top", top), ("bottom", bottom), ("left", left), ("right", right),
    ]:
        if val is None:
            continue
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} '
            f'w:val="single" '
            f'w:sz="{int(val["width"] * 8)}" '
            f'w:color="{val["color"]}" '
            f'w:space="0"/>'
        )
        existing = tcBorders.find(qn(f"w:{edge}"))
        if existing is not None:
            tcBorders.remove(existing)
        tcBorders.append(element)


def _set_cell_padding(cell, top_pt=0, bottom_pt=0, left_pt=0, right_pt=0):
    """Set cell padding (margins) via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = parse_xml(f'<w:tcMar {nsdecls("w")} />')
        tcPr.append(tcMar)
    else:
        tcMar.clear()

    for edge, pt_val in [
        ("top", top_pt), ("bottom", bottom_pt),
        ("start", left_pt), ("end", right_pt),
    ]:
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} '
            f'w:w="{int(pt_val * 20)}" w:type="dxa"/>'
        )
        tcMar.append(el)


def _set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shading)


# ============================================================
# Cell styling helpers
# ============================================================


def _style_cell_text(
    cell,
    text: str,
    bold: bool = False,
    font_size: Pt | None = None,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
):
    """Set text content and formatting for a cell."""
    cell.text = ""
    para = cell.paragraphs[0]
    if alignment:
        para.alignment = alignment
    run = para.add_run(text)
    run.font.name = FONT_FAMILY
    run.font.size = font_size or FONT_SIZE_DEFAULT
    if bold:
        run.font.bold = True


def _apply_default_borders(cell):
    """Apply standard border styling to a cell."""
    spec = {"width": BORDER_WIDTH_PT, "color": BORDER_COLOR_HEX}
    _set_cell_border(cell, top=spec, bottom=spec, left=spec, right=spec)


def _apply_default_padding(cell):
    """Apply standard cell padding."""
    _set_cell_padding(
        cell,
        CELL_PAD_TOP_PT, CELL_PAD_BOTTOM_PT,
        CELL_PAD_LEFT_PT, CELL_PAD_RIGHT_PT,
    )


def _set_paragraph_spacing(cell):
    """Set paragraph spacing within a cell."""
    for para in cell.paragraphs:
        pf = para.paragraph_format
        pf.space_before = Pt(CELL_SPACE_ABOVE_PT)
        pf.space_after = Pt(CELL_SPACE_BELOW_PT)


def _style_table(table: Table, col_widths_pt: list[float], font_size: Pt | None = None):
    """Apply full styling to a table: borders, padding, widths, header row."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.width = _pt_to_emu(col_widths_pt[col_idx])
            _apply_default_borders(cell)
            _apply_default_padding(cell)
            _set_paragraph_spacing(cell)

            if row_idx == 0:
                _set_cell_shading(cell, HEADER_BG_HEX)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.bold = True
                        run.font.name = FONT_FAMILY
                        run.font.size = font_size or FONT_SIZE_DEFAULT
            else:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = FONT_FAMILY
                        run.font.size = font_size or FONT_SIZE_DEFAULT


# ============================================================
# Public table builders
# ============================================================


def build_info_table(doc: Document, label: str, description: str) -> Table:
    """Build a 2×2 Info Table for a component section."""
    table = doc.add_table(rows=2, cols=2)
    _style_cell_text(
        table.cell(0, 0), "Tên chức năng",
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _style_cell_text(
        table.cell(0, 1), f"Component {label}",
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _style_cell_text(table.cell(1, 0), "Mô tả")
    _style_cell_text(table.cell(1, 1), description)
    _style_table(table, INFO_COLS_PT)
    return table


def build_screen_level_info_table(
    doc: Document, screen_name: str, description: str,
) -> Table:
    """Build a 2×2 Info Table for the screen overview section."""
    table = doc.add_table(rows=2, cols=2)
    _style_cell_text(
        table.cell(0, 0), "Tên màn hình",
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _style_cell_text(
        table.cell(0, 1), f"Màn hình {screen_name}",
        bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _style_cell_text(table.cell(1, 0), "Mô tả")
    _style_cell_text(table.cell(1, 1), description)
    _style_table(table, INFO_COLS_PT)
    return table


def build_ui_elements_table(
    doc: Document, children: list[ChildElement],
) -> Table:
    """Build a 7-column UI Elements Table."""
    table = doc.add_table(rows=1 + len(children), cols=7)

    # Header
    for c, text in enumerate(UI_TABLE_HEADERS):
        _style_cell_text(
            table.cell(0, c), text,
            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    # Data
    for r, child in enumerate(children):
        row_data = [
            str(child.stt), child.label, child.controlType,
            child.required, child.maxLength, child.editable,
            child.description,
        ]
        for c, text in enumerate(row_data):
            _style_cell_text(table.cell(r + 1, c), text)

    _style_table(table, UI_COLS_PT)
    return table


def build_interaction_table(
    doc: Document, interactions: list[Interaction],
) -> Table:
    """Build a 2-column Interaction Events Table."""
    table = doc.add_table(rows=1 + len(interactions), cols=2)

    for c, text in enumerate(INTERACTION_TABLE_HEADERS):
        _style_cell_text(
            table.cell(0, c), text,
            bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    for r, interaction in enumerate(interactions):
        _style_cell_text(table.cell(r + 1, 0), interaction.action)
        _style_cell_text(table.cell(r + 1, 1), interaction.reaction)

    _style_table(table, INTERACTION_COLS_PT)
    return table


def build_api_table(doc: Document, params: list[ApiParam]) -> Table:
    """Build a 6-column API Parameter Table (for request or response)."""
    table = doc.add_table(rows=1 + len(params), cols=6)

    for c, text in enumerate(API_TABLE_HEADERS):
        _style_cell_text(
            table.cell(0, c), text,
            bold=True, font_size=FONT_SIZE_API,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    for r, param in enumerate(params):
        row_data = [
            param.name, param.meaning, param.required,
            param.dataType, param.limit, param.defaultValue,
        ]
        for c, text in enumerate(row_data):
            _style_cell_text(table.cell(r + 1, c), text, font_size=FONT_SIZE_API)

    _style_table(table, API_COLS_PT, font_size=FONT_SIZE_API)
    return table
