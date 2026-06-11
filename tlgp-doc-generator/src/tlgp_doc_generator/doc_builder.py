"""Orchestrate the full .docx build from analysis.json data."""

from __future__ import annotations

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt

from tlgp_doc_generator.image_handler import insert_image
from tlgp_doc_generator.models import AnalysisData, Api, Component, Screen
from tlgp_doc_generator.style_constants import (
    FONT_FAMILY,
    FONT_SIZE_DEFAULT,
    HEADING_COLOR,
)
from tlgp_doc_generator.table_builder import (
    build_api_table,
    build_info_table,
    build_interaction_table,
    build_screen_level_info_table,
    build_ui_elements_table,
)


# ============================================================
# Heading helpers
# ============================================================


def _set_run_font(run, font_name: str = FONT_FAMILY):
    """Ensure both Latin and East Asian fonts are set."""
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _add_h3(doc: Document, text: str):
    """Component title: HEADING_3, bold, #4F81BD."""
    para = doc.add_heading(level=3)
    run = para.add_run(text)
    _set_run_font(run)
    run.font.size = FONT_SIZE_DEFAULT
    run.font.color.rgb = HEADING_COLOR
    run.font.bold = True


def _add_h4(doc: Document, text: str):
    """Sub-section title: HEADING_4, italic, #4F81BD."""
    para = doc.add_heading(level=4)
    run = para.add_run(text)
    _set_run_font(run)
    run.font.size = FONT_SIZE_DEFAULT
    run.font.color.rgb = HEADING_COLOR
    run.font.italic = True
    run.font.bold = False


def _add_bold_text(doc: Document, text: str):
    """Bold normal text (for API headings, request/response labels)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run)
    run.font.size = FONT_SIZE_DEFAULT
    run.font.bold = True


def _add_normal_text(doc: Document, text: str):
    """Plain normal text (for API URLs, sub-DTO titles)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run)
    run.font.size = FONT_SIZE_DEFAULT


# ============================================================
# Section builders
# ============================================================


def _add_component_section(
    doc: Document,
    component: Component,
    section_number: str,
    analysis: AnalysisData,
):
    """Build a complete component section with all sub-sections."""
    # H3: Component title
    _add_h3(doc, f"{section_number}. Component {component.label}")

    # H4: 1. General info
    _add_h4(doc, f"{section_number}.1 Thông tin chung về chức năng")
    build_info_table(doc, component.label, component.description)

    # H4: 2. Screen image
    _add_h4(doc, f"{section_number}.2 Màn hình chức năng")
    if component.imageFile:
        image_path = analysis.resolve_image(component.imageFile)
        insert_image(doc, image_path, full_width=True)

    # H4: 3. UI elements
    if component.children:
        _add_h4(doc, f"{section_number}.3 Mô tả chi tiết các thành phần trên màn hình")
        build_ui_elements_table(doc, component.children)

    # H4: 4. Interaction events
    if component.interactions:
        step_count = len(component.interactions)
        _add_h4(
            doc,
            f"{section_number}.4 Xử lý luồng sự kiện tương tác ({step_count} bước)",
        )
        build_interaction_table(doc, component.interactions)


def _add_screen_section(
    doc: Document,
    screen: Screen,
    section_number: str,
    analysis: AnalysisData,
):
    """Build the screen overview section."""
    # H3: Screen title
    _add_h3(doc, f"{section_number}. Màn hình {screen.name}")

    # H4: 1. General info
    _add_h4(doc, f"{section_number}.1 Thông tin chung về chức năng")
    build_screen_level_info_table(doc, screen.name, screen.description)

    # H4: 2. Screen image(s)
    _add_h4(doc, f"{section_number}.2 Màn hình chức năng")
    for img_file in screen.imageFiles:
        image_path = analysis.resolve_image(img_file)
        insert_image(doc, image_path, full_width=True)

    # H4: 3. UI elements (top-level children)
    if screen.topLevelChildren:
        _add_h4(doc, f"{section_number}.3 Mô tả chi tiết các thành phần trên màn hình")
        build_ui_elements_table(doc, screen.topLevelChildren)

    # H4: 4. Interaction events
    if screen.interactions:
        step_count = len(screen.interactions)
        _add_h4(
            doc,
            f"{section_number}.4 Xử lý luồng sự kiện tương tác ({step_count} bước)",
        )
        build_interaction_table(doc, screen.interactions)


def _add_api_section(doc: Document, api: Api):
    """Build a single API documentation block."""
    # API title: bold normal text
    _add_bold_text(doc, f"{api.number}. {api.method} {api.title}")

    # API URL: plain normal text
    _add_normal_text(doc, f"URL: {api.url}")

    # Request section
    if api.requestParams:
        if api.requestBodyType:
            _add_bold_text(doc, f"Request Body ({api.requestBodyType})")
        else:
            _add_bold_text(doc, "Request")
        build_api_table(doc, api.requestParams)
    elif api.requestDescription:
        _add_bold_text(doc, "Request")
        _add_normal_text(doc, api.requestDescription)

    # Response section
    if api.responseFields:
        response_label = (
            f"Response (data = {api.responseType})" if api.responseType
            else "Response"
        )
        _add_bold_text(doc, response_label)
        build_api_table(doc, api.responseFields)
    elif api.responseType:
        _add_bold_text(doc, f"Response (data = {api.responseType})")
        if api.responseDescription:
            _add_normal_text(doc, api.responseDescription)

    # Sub-DTO tables
    for sub in api.subDtos:
        title = f"{sub.name} fields"
        if sub.fieldRef:
            title += f" ({sub.fieldRef})"
        _add_normal_text(doc, title)
        if sub.fields:
            build_api_table(doc, sub.fields)


# ============================================================
# Public API
# ============================================================


def build_document(analysis: AnalysisData) -> Document:
    """Build a complete .docx from analysis data.

    Follows the TLGP spec structure:
    1. Non-leaf component sections (in annotation order)
    2. Screen overview section
    3. API documentation
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = FONT_FAMILY
    style.font.size = FONT_SIZE_DEFAULT

    # Set page margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Collect non-leaf components for section numbering
    non_leaf_components = [c for c in analysis.components if not c.isLeaf]

    # Component sections
    for i, component in enumerate(non_leaf_components):
        section_num = f"{analysis.sectionPrefix}.{i + 1}"
        _add_component_section(doc, component, section_num, analysis)

    # Screen overview section
    screen_section_num = f"{analysis.sectionPrefix}.{len(non_leaf_components) + 1}"
    _add_screen_section(doc, analysis.screen, screen_section_num, analysis)

    # API documentation
    for api in analysis.apis:
        _add_api_section(doc, api)

    return doc
