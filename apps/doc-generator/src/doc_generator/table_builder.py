"""Build styled tables for each TLGP table type.

Uses python-docx with XML manipulation for properties not exposed
by the high-level API (borders, cell padding, shading).
"""

from __future__ import annotations

from collections.abc import Sequence

from docx.document import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt
from docx.table import Table

from doc_generator.models import (
    AnalysisData,
    ApiParam,
    ChildElement,
    Interaction,
)
from doc_generator.style_constants import StyleConfig

# ============================================================
# Low-level XML helpers
# ============================================================


def _pt_to_twips(pt: float) -> int:
    """Convert points to twips (twentieths of a point / DXA).

    OOXML uses twips as the native unit for table and cell widths.
    1 pt = 20 twips.
    """
    return int(pt * 20)


def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set cell borders via XML."""
    tc = cell._tc  # noqa: SLF001
    tcPr = tc.get_or_add_tcPr()

    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = parse_xml(f"<w:tcBorders {nsdecls('w')} />")
        tcPr.append(tcBorders)

    for edge, val in [
        ("top", top),
        ("bottom", bottom),
        ("left", left),
        ("right", right),
    ]:
        if val is None:
            continue
        element = parse_xml(
            f"<w:{edge} {nsdecls('w')} "
            f'w:val="single" '
            f'w:sz="{int(val["width"] * 8)}" '
            f'w:color="{val["color"]}" '
            f'w:space="0"/>'
        )
        existing = tcBorders.find(qn(f"w:{edge}"))
        if existing is not None:
            tcBorders.remove(existing)
        tcBorders.append(element)


def _set_cell_padding(cell, top_pt=0, bottom_pt=0, left_pt=0, right_pt=0):
    """Set cell padding (margins) via XML."""
    tc = cell._tc  # noqa: SLF001
    tcPr = tc.get_or_add_tcPr()

    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = parse_xml(f"<w:tcMar {nsdecls('w')} />")
        tcPr.append(tcMar)
    else:
        tcMar.clear()

    for edge, pt_val in [
        ("top", top_pt),
        ("bottom", bottom_pt),
        ("start", left_pt),
        ("end", right_pt),
    ]:
        el = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:w="{int(pt_val * 20)}" w:type="dxa"/>'
        )
        tcMar.append(el)


def _set_cell_shading(cell, color_hex: str):
    """Set cell background color."""
    tc = cell._tc  # noqa: SLF001
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shading)


def _set_cell_width(cell, twips: int):
    """Set cell width directly in twips (DXA) via XML.

    Bypasses python-docx's EMU-based cell.width property for exact
    twips values that match tblGrid gridCol declarations.
    """
    tc = cell._tc  # noqa: SLF001
    tcPr = tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is not None:
        tcPr.remove(tcW)
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="{twips}" w:type="dxa"/>')
    # Insert tcW as the first child of tcPr for schema compliance
    tcPr.insert(0, tcW)


# ============================================================
# Cell styling helpers
# ============================================================


def _style_cell_text(
    cell,
    text: str,
    style: StyleConfig,
    bold: bool = False,
    font_size: Pt | None = None,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
):
    """Set text content and formatting for a cell.

    Uses cell.text to produce a single clean run, then applies formatting
    to that run.
    """
    cell.text = text
    para = cell.paragraphs[0]
    if alignment:
        para.alignment = alignment
    for run in para.runs:
        run.font.name = style.FONT_FAMILY
        run.font.size = font_size or style.FONT_SIZE_DEFAULT
        if bold:
            run.font.bold = True


def _apply_default_borders(cell, style: StyleConfig):
    """Apply standard border styling to a cell."""
    spec = {"width": style.BORDER_WIDTH_PT, "color": style.BORDER_COLOR_HEX}
    _set_cell_border(cell, top=spec, bottom=spec, left=spec, right=spec)


def _apply_default_padding(cell, style: StyleConfig):
    """Apply standard cell padding."""
    _set_cell_padding(
        cell,
        style.CELL_PAD_TOP_PT,
        style.CELL_PAD_BOTTOM_PT,
        style.CELL_PAD_LEFT_PT,
        style.CELL_PAD_RIGHT_PT,
    )


def _set_paragraph_spacing(cell, style: StyleConfig):
    """Set paragraph spacing within a cell."""
    for para in cell.paragraphs:
        pf = para.paragraph_format
        pf.space_before = Pt(style.CELL_SPACE_ABOVE_PT)
        pf.space_after = Pt(style.CELL_SPACE_BELOW_PT)


def _set_fixed_table_layout(table: Table, col_widths_pt: list[float]):
    """Enforce a fixed-width table layout with exact column widths.

    Sets three properties that together guarantee column widths render
    identically across Word, LibreOffice, and Google Docs:
      - tblW:      total width in twips, type=dxa (not auto)
      - tblLayout: type=fixed (disables autofit)
      - gridCol:   per-column widths matching tcW values
    """
    col_twips = [_pt_to_twips(pt) for pt in col_widths_pt]
    total_twips = sum(col_twips)

    tbl = table._tbl  # noqa: SLF001
    tblPr = tbl.find(qn("w:tblPr"))

    # Set total table width
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is not None:
        tblW.set(qn("w:w"), str(total_twips))
        tblW.set(qn("w:type"), "dxa")
    else:
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{total_twips}" w:type="dxa"/>')
        tblPr.insert(0, tblW)

    # Disable autofit
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is not None:
        tblLayout.set(qn("w:type"), "fixed")
    else:
        tblLayout = parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>')
        tblPr.append(tblLayout)

    # Replace gridCol elements with correct per-column widths
    tblGrid = tbl.find(qn("w:tblGrid"))
    for gc in tblGrid.findall(qn("w:gridCol")):
        tblGrid.remove(gc)
    for twips in col_twips:
        gc = parse_xml(f'<w:gridCol {nsdecls("w")} w:w="{twips}"/>')
        tblGrid.append(gc)


def _style_table(
    table: Table,
    col_widths_pt: list[float],
    style: StyleConfig,
    font_size: Pt | None = None,
):
    """Apply full styling to a table: layout, borders, padding, widths, header row."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_fixed_table_layout(table, col_widths_pt)

    col_twips = [_pt_to_twips(pt) for pt in col_widths_pt]

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            _set_cell_width(cell, col_twips[col_idx])
            _apply_default_borders(cell, style)
            _apply_default_padding(cell, style)
            _set_paragraph_spacing(cell, style)

            if row_idx == 0:
                _set_cell_shading(cell, style.HEADER_BG_HEX)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.bold = True
                        run.font.name = style.FONT_FAMILY
                        run.font.size = font_size or style.FONT_SIZE_DEFAULT
            else:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = style.FONT_FAMILY
                        run.font.size = font_size or style.FONT_SIZE_DEFAULT


# ============================================================
# Public table builders
# ============================================================


def build_generic_info_table(
    doc: Document,
    header_label: str,
    header_value: str,
    description: str,
    style: StyleConfig,
) -> Table:
    """Build a 2×2 Info Table for a section."""
    table = doc.add_table(rows=2, cols=2)
    _style_cell_text(
        table.cell(0, 0),
        header_label,
        style,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _style_cell_text(
        table.cell(0, 1),
        header_value,
        style,
        bold=True,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    _style_cell_text(table.cell(1, 0), "Mô tả", style)
    _style_cell_text(table.cell(1, 1), description, style)
    _style_table(table, style.INFO_COLS_PT, style)
    return table


def build_ui_elements_table(
    doc: Document,
    children: Sequence[ChildElement],
    style: StyleConfig,
    analysis: AnalysisData | None = None,
) -> Table:
    """Build a 7-column UI Elements Table."""
    table = doc.add_table(rows=1 + len(children), cols=7)

    # Header
    for c, text in enumerate(style.UI_TABLE_HEADERS):
        _style_cell_text(
            table.cell(0, c),
            text,
            style,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    # Data
    for r, child in enumerate(children):
        label = child.label
        description = child.description

        if child.type == "component" and analysis is not None:
            target = analysis.components.get(child.componentId)
            if target:
                label = label or target.label
                description = description or target.description

        row_data = [
            str(r + 1),
            label,
            child.controlType,
            getattr(child, "required", ""),
            getattr(child, "maxLength", ""),
            getattr(child, "editable", ""),
            description,
        ]
        for c, text in enumerate(row_data):
            _style_cell_text(table.cell(r + 1, c), text, style)

    _style_table(table, style.UI_COLS_PT, style)
    return table


def build_interaction_table(
    doc: Document,
    interactions: list[Interaction],
    style: StyleConfig,
) -> Table:
    """Build a 2-column Interaction Events Table."""
    table = doc.add_table(rows=1 + len(interactions), cols=2)

    for c, text in enumerate(style.INTERACTION_TABLE_HEADERS):
        _style_cell_text(
            table.cell(0, c),
            text,
            style,
            bold=True,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    for r, interaction in enumerate(interactions):
        _style_cell_text(table.cell(r + 1, 0), interaction.action, style)
        _style_cell_text(table.cell(r + 1, 1), interaction.reaction, style)

    _style_table(table, style.INTERACTION_COLS_PT, style)
    return table


def build_api_table(doc: Document, params: list[ApiParam], style: StyleConfig) -> Table:
    """Build a 6-column API Parameter Table (for request or response)."""
    table = doc.add_table(rows=1 + len(params), cols=6)

    for c, text in enumerate(style.API_TABLE_HEADERS):
        _style_cell_text(
            table.cell(0, c),
            text,
            style,
            bold=True,
            font_size=style.FONT_SIZE_API,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
        )

    for r, param in enumerate(params):
        row_data = [
            param.name,
            param.meaning,
            param.required,
            param.dataType,
            param.limit,
            param.defaultValue,
        ]
        for c, text in enumerate(row_data):
            _style_cell_text(
                table.cell(r + 1, c), text, style, font_size=style.FONT_SIZE_API
            )

    _style_table(table, style.API_COLS_PT, style, font_size=style.FONT_SIZE_API)
    return table
