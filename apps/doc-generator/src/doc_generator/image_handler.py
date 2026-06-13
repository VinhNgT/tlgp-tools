"""Image embedding for annotated screenshots."""

from __future__ import annotations

from pathlib import Path

from docx.document import Document
from docx.shared import Inches

# Full page width matches the table width (6.5 inches = 468pt)
IMAGE_FULL_WIDTH = Inches(6.5)


def insert_image(
    doc: Document,
    image_path: Path,
    full_width: bool = True,
) -> None:
    """Embed an image into the document.

    Args:
        doc: The document to insert into.
        image_path: Absolute path to the image file.
        full_width: If True, scale to full page width (468pt).
                    If False, use natural size.
    """
    if not image_path.exists():
        # Insert a placeholder paragraph if image is missing
        para = doc.add_paragraph()
        run = para.add_run(f"[Image not found: {image_path.name}]")
        run.font.italic = True
        return

    width = IMAGE_FULL_WIDTH if full_width else None
    doc.add_picture(str(image_path), width=width)
