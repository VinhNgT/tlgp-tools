"""Orchestrate the full .docx build from analysis.json data."""

from __future__ import annotations

from typing import Any

import docx
from docx.document import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt
from tlgp_contracts import Api, ApiPayload, NodeSpec, ScreenSpec

from doc_generator.image_handler import insert_image
from doc_generator.style_constants import StyleConfig, load_default_style
from doc_generator.table_builder import (
    build_api_table,
    build_generic_info_table,
    build_interaction_table,
    build_ui_elements_table,
)

# ============================================================
# Heading helpers
# ============================================================


def _set_run_font(run, style: StyleConfig):
    """Ensure both Latin and East Asian fonts are set."""
    font_name = style.FONT_FAMILY
    run.font.name = font_name
    # Access the underlying XML element to configure the East Asian font property,
    # as python-docx does not provide a public high-level API for this setting.
    rPr = run._element.get_or_add_rPr()  # noqa: SLF001
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f"<w:rFonts {nsdecls('w')}/>")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _add_h3(doc: Document, text: str, style: StyleConfig):
    """Component title: HEADING_3, bold, #4F81BD."""
    para = doc.add_heading(level=3)
    run = para.add_run(text)
    _set_run_font(run, style)
    run.font.size = style.FONT_SIZE_DEFAULT
    run.font.color.rgb = style.HEADING_COLOR
    run.font.bold = True


def _add_h4(doc: Document, text: str, style: StyleConfig):
    """Sub-section title: HEADING_4, italic, #4F81BD."""
    para = doc.add_heading(level=4)
    run = para.add_run(text)
    _set_run_font(run, style)
    run.font.size = style.FONT_SIZE_DEFAULT
    run.font.color.rgb = style.HEADING_COLOR
    run.font.italic = True
    run.font.bold = False


def _add_bold_text(doc: Document, text: str, style: StyleConfig):
    """Bold normal text (for API headings, request/response labels)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, style)
    run.font.size = style.FONT_SIZE_DEFAULT
    run.font.bold = True


def _add_normal_text(doc: Document, text: str, style: StyleConfig):
    """Plain normal text (for API URLs, sub-DTO titles)."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    _set_run_font(run, style)
    run.font.size = style.FONT_SIZE_DEFAULT


# ============================================================
# Section builders
# ============================================================


def _add_content_section(
    doc: Document,
    title: str,
    section_number: str,
    info_label: str,
    info_value: str,
    info_description: str,
    image_files: list[str],
    children: list,
    interactions: list,
    apis: list,
    analysis: ScreenSpec,
    style: StyleConfig,
):
    """Build a complete section with all sub-sections (used for both components and screens)."""
    # H3: Title
    _add_h3(doc, f"{section_number}. {title}", style)

    sub = 1

    # H4: 1. General info
    _add_h4(doc, f"{section_number}.{sub} Thông tin chung về chức năng", style)
    build_generic_info_table(doc, info_label, info_value, info_description, style)
    sub += 1

    # H4: 2. Screen image
    _add_h4(doc, f"{section_number}.{sub} Màn hình chức năng", style)
    for img_file in image_files:
        if img_file:
            image_path = analysis.resolve_annotated_image(img_file)
            insert_image(doc, image_path, full_width=True)
    sub += 1

    # H4: 3. UI elements
    _add_h4(
        doc,
        f"{section_number}.{sub} Mô tả chi tiết các thành phần trên màn hình",
        style,
    )
    if children:
        build_ui_elements_table(doc, children, style, analysis=analysis)
    sub += 1

    # H4: 4. Interaction events
    if interactions:
        step_count = len(interactions)
        _add_h4(
            doc,
            f"{section_number}.{sub} Xử lý luồng sự kiện tương tác ({step_count} bước)",
            style,
        )
        build_interaction_table(doc, interactions, style)
        sub += 1

    # Render APIs at the end of the section
    for idx, api in enumerate(apis):
        _add_api_section(doc, api, style, idx + 1)


def _add_component_section(
    doc: Document,
    component: NodeSpec,
    section_number: str,
    analysis: ScreenSpec,
    style: StyleConfig,
):
    """Build a complete component section with all sub-sections."""
    _add_content_section(
        doc,
        title=f"Component {component.label}",
        section_number=section_number,
        info_label="Tên chức năng",
        info_value=f"Component {component.label}",
        info_description=component.description,
        image_files=component.annotatedImages,
        children=component.childrenIds,
        interactions=component.interactions,
        apis=component.apis,
        analysis=analysis,
        style=style,
    )


def _add_screen_section(
    doc: Document,
    screen: NodeSpec,
    section_number: str,
    analysis: ScreenSpec,
    style: StyleConfig,
):
    """Build the screen overview section."""
    _add_content_section(
        doc,
        title=f"Màn hình {screen.label}",
        section_number=section_number,
        info_label="Tên màn hình",
        info_value=f"Màn hình {screen.label}",
        info_description=screen.description,
        image_files=screen.annotatedImages,
        children=screen.childrenIds,
        interactions=screen.interactions,
        apis=screen.apis,
        analysis=analysis,
        style=style,
    )


def _add_payload_section(
    doc: Document,
    root_type: str | None,
    payloads: list[ApiPayload],
    is_response: bool,
    style: StyleConfig,
):
    """Render request or response payloads list starting from the root type."""
    if not root_type:
        heading = "Response" if is_response else "Request Body"
        _add_bold_text(doc, heading, style)
        text = "Không có dữ liệu trả về" if is_response else "Không có tham số"
        _add_normal_text(doc, text, style)
        return

    if is_response:
        heading = f"Response (data = {root_type})"
    else:
        heading = f"Request Body ({root_type})"
    _add_bold_text(doc, heading, style)

    # Sort payloads so that the root payload is first, preserving the order of others
    root_payload = None
    other_payloads = []
    for p in payloads:
        if p.type == root_type:
            root_payload = p
        else:
            other_payloads.append(p)

    ordered_payloads = []
    if root_payload:
        ordered_payloads.append(root_payload)
    ordered_payloads.extend(other_payloads)

    for idx, payload in enumerate(ordered_payloads):
        # Print DTO type heading only if it's a child DTO (i.e., not the root_type)
        if idx > 0 or payload.type != root_type:
            # Child DTO heading
            _add_normal_text(doc, payload.type, style)

        if payload.fields:
            build_api_table(doc, payload.fields, style)


def _add_api_section(doc: Document, api: Api, style: StyleConfig, api_index: int):
    """Build a single API documentation block."""
    # API title: bold normal text
    para = doc.add_paragraph()
    run = para.add_run(f"{api_index}. {api.name}")
    _set_run_font(run, style)
    run.font.size = style.FONT_SIZE_DEFAULT
    run.font.bold = True
    para.paragraph_format.space_before = Pt(style.API_SECTION_SPACE_BEFORE_PT)

    # API URL: plain normal text
    _add_normal_text(doc, f"URL: {api.url}", style)

    # 1. Request Payload
    _add_payload_section(
        doc,
        api.requestRootType,
        api.request,
        is_response=False,
        style=style,
    )

    # 2. Response Payload
    _add_payload_section(
        doc,
        api.responseRootType,
        api.response,
        is_response=True,
        style=style,
    )


# ============================================================
# Style configuration
# ============================================================


def _configure_document_styles(doc: Document, style: StyleConfig):
    """Configure paragraph and heading styles for consistent vertical rhythm.

    Spacing is applied at the style level so the document uses proper
    Word style definitions — the same approach a professional would use
    when creating a document by hand.
    """
    # Normal style: font and paragraph spacing
    normal_style: Any = doc.styles["Normal"]
    normal_style.font.name = style.FONT_FAMILY
    normal_style.font.size = style.FONT_SIZE_DEFAULT
    normal_style.paragraph_format.space_before = Pt(style.NORMAL_SPACE_BEFORE_PT)
    normal_style.paragraph_format.space_after = Pt(style.NORMAL_SPACE_AFTER_PT)

    # Heading 3: section titles (e.g., "1.1.1. Component X")
    h3_style: Any = doc.styles["Heading 3"]
    h3_style.paragraph_format.space_before = Pt(style.H3_SPACE_BEFORE_PT)
    h3_style.paragraph_format.space_after = Pt(style.H3_SPACE_AFTER_PT)

    # Heading 4: sub-section titles (e.g., "1.1.1.1 Thông tin chung")
    h4_style: Any = doc.styles["Heading 4"]
    h4_style.paragraph_format.space_before = Pt(style.H4_SPACE_BEFORE_PT)
    h4_style.paragraph_format.space_after = Pt(style.H4_SPACE_AFTER_PT)


# ============================================================
# Public API
# ============================================================


def build_document(analysis: ScreenSpec) -> Document:
    """Build a complete .docx from spec data.

    Follows the TLGP spec structure:
    1. Sub-component sections (bottom-up order)
    2. Screen overview section
    """
    doc = docx.Document()
    style_config = load_default_style()

    # Configure all document styles
    _configure_document_styles(doc, style_config)

    # Set page margins (1 inch all around)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Find all components using post-order DFS starting from rootId
    dfs_order: list[int] = []
    visited: set[int] = set()
    nodes_dict = analysis.nodes_map

    def dfs(node_id: int):
        if node_id in visited:
            return
        visited.add(node_id)
        node = nodes_dict.get(node_id)
        if not node:
            return
        for child_id in node.childrenIds:
            child = nodes_dict.get(child_id)
            # A node is a sub-component if it has children of its own
            if child and len(child.childrenIds) > 0:
                dfs(child_id)
        if node_id != analysis.rootId:
            dfs_order.append(node_id)

    dfs(analysis.rootId)
    sub_components = [nodes_dict[cid] for cid in dfs_order if cid in nodes_dict]

    # Component sections
    for i, component in enumerate(sub_components):
        section_num = f"{analysis.sectionPrefix}.{i + 1}"
        _add_component_section(doc, component, section_num, analysis, style_config)

    # Screen overview section (rendered last)
    screen_section_num = f"{analysis.sectionPrefix}.{len(sub_components) + 1}"
    _add_screen_section(
        doc, analysis.screen, screen_section_num, analysis, style_config
    )

    return doc
