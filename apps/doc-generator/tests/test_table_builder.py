"""Tests for table_builder — verifies table structure and styling."""

import pytest
from tlgp_contracts import (
    ApiParam,
    Interaction,
    NodeSpec,
    ScreenSpec,
)
from doc_generator.style_constants import StyleConfig, load_default_style
from doc_generator.table_builder import (
    build_api_table,
    build_generic_info_table,
    build_interaction_table,
    build_ui_elements_table,
)
from docx import Document
from docx.oxml.ns import qn

style: StyleConfig


@pytest.fixture(autouse=True, scope="module")
def setup_style():
    global style
    style = load_default_style()


def _get_cell_shading(cell) -> str | None:
    """Extract the fill color from a cell's shading element."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return None
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        return None
    return shd.get(qn("w:fill"))


def _get_cell_border_color(cell, edge: str) -> str | None:
    """Extract a border color from a cell."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        return None
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        return None
    el = borders.find(qn(f"w:{edge}"))
    if el is None:
        return None
    return el.get(qn("w:color"))


# ── Info Table ────────────────────────────────────────────────────────


class TestGenericInfoTable:
    def test_dimensions(self):
        doc = Document()
        table = build_generic_info_table(doc, "L", "V", "Description text", style)
        assert len(table.rows) == 2
        assert len(table.columns) == 2

    def test_header_content(self):
        doc = Document()
        table = build_generic_info_table(
            doc, "Tên chức năng", "Tiêu đề", "Mô tả...", style
        )
        assert table.cell(0, 0).text == "Tên chức năng"
        assert "Tiêu đề" in table.cell(0, 1).text

    def test_data_content(self):
        doc = Document()
        table = build_generic_info_table(doc, "X", "Y", "Detailed description", style)
        assert table.cell(1, 0).text == "Mô tả"
        assert table.cell(1, 1).text == "Detailed description"

    def test_header_row_has_background(self):
        doc = Document()
        table = build_generic_info_table(doc, "A", "B", "C", style)
        # Both header cells should have background
        assert _get_cell_shading(table.cell(0, 0)) == style.HEADER_BG_HEX
        assert _get_cell_shading(table.cell(0, 1)) == style.HEADER_BG_HEX

    def test_borders_applied(self):
        doc = Document()
        table = build_generic_info_table(doc, "A", "B", "C", style)
        for row in table.rows:
            for cell in row.cells:
                assert _get_cell_border_color(cell, "top") == style.BORDER_COLOR_HEX
                assert _get_cell_border_color(cell, "bottom") == style.BORDER_COLOR_HEX


# ── UI Elements Table ─────────────────────────────────────────────────


class TestUiElementsTable:
    def test_correct_dimensions(self):
        doc = Document()
        analysis = ScreenSpec(
            imageDir=".",
            nodes=[
                NodeSpec(id="1", label="Back", controlType="Icon"),
                NodeSpec(id="2", label="Title", controlType="Text"),
            ]
        )
        table = build_ui_elements_table(doc, ["1", "2"], style, analysis=analysis)
        assert len(table.rows) == 3  # 1 header + 2 data
        assert len(table.columns) == 7

    def test_header_labels(self):
        doc = Document()
        analysis = ScreenSpec(
            imageDir=".",
            nodes=[
                NodeSpec(id="1", label="A", controlType="B"),
            ]
        )
        table = build_ui_elements_table(doc, ["1"], style, analysis=analysis)
        headers = [table.cell(0, c).text for c in range(7)]
        assert headers[0] == "STT"
        assert headers[1] == "Tên"
        assert headers[6] == "Mô tả"

    def test_data_populated(self):
        doc = Document()
        analysis = ScreenSpec(
            imageDir=".",
            nodes=[
                NodeSpec(
                    id=1,
                    label="Share",
                    controlType="Icon",
                    description="Share product",
                    required=True,
                    editable=False,
                    maxLength=20,
                ),
            ]
        )
        table = build_ui_elements_table(doc, [1], style, analysis=analysis)
        assert table.cell(1, 0).text == "1"
        assert table.cell(1, 1).text == "Share"
        assert table.cell(1, 2).text == "Icon"
        assert table.cell(1, 3).text == "Có"
        assert table.cell(1, 4).text == "20"
        assert table.cell(1, 5).text == "Không"
        assert table.cell(1, 6).text == "Share product"

    def test_empty_children_produces_header_only(self):
        doc = Document()
        table = build_ui_elements_table(doc, [], style)
        assert len(table.rows) == 1  # header only


# ── Interaction Table ─────────────────────────────────────────────────


class TestInteractionTable:
    def test_correct_dimensions(self):
        doc = Document()
        interactions = [
            Interaction(action="Click Back", reaction="Go back"),
        ]
        table = build_interaction_table(doc, interactions, style)
        assert len(table.rows) == 2  # 1 header + 1 data
        assert len(table.columns) == 2

    def test_header_labels(self):
        doc = Document()
        interactions = [Interaction(action="A", reaction="B")]
        table = build_interaction_table(doc, interactions, style)
        assert table.cell(0, 0).text == "Hành động của tác nhân"
        assert table.cell(0, 1).text == "Phản ứng của hệ thống"

    def test_data_populated(self):
        doc = Document()
        interactions = [
            Interaction(action="Tap share", reaction="Open share sheet"),
        ]
        table = build_interaction_table(doc, interactions, style)
        assert table.cell(1, 0).text == "Tap share"
        assert table.cell(1, 1).text == "Open share sheet"


# ── API Table ─────────────────────────────────────────────────────────


class TestApiTable:
    def test_correct_dimensions(self):
        doc = Document()
        params = [
            ApiParam(name="id", description="ID", type="String"),
        ]
        table = build_api_table(doc, params, style)
        assert table.cell(1, 0).text == "id"
        assert len(table.rows) == 2
        assert len(table.columns) == 6

    def test_header_labels(self):
        doc = Document()
        params = [ApiParam(name="x")]
        table = build_api_table(doc, params, style)
        headers = [table.cell(0, c).text for c in range(6)]
        assert headers[0] == "Tên tham số"
        assert headers[3] == "Kiểu dữ liệu"

    def test_data_populated(self):
        doc = Document()
        params = [
            ApiParam(
                name="merchant_id",
                description="Mã đối tác",
                required="Có",
                type="String",
            ),
        ]
        table = build_api_table(doc, params, style)
        assert table.cell(1, 0).text == "merchant_id"
        assert table.cell(1, 1).text == "Mã đối tác"
        assert table.cell(1, 2).text == "Có"
        assert table.cell(1, 3).text == "String"

    def test_multiple_params(self):
        doc = Document()
        params = [
            ApiParam(name="a"),
            ApiParam(name="b"),
            ApiParam(name="c"),
        ]
        table = build_api_table(doc, params, style)
        assert len(table.rows) == 4  # 1 header + 3 data

    def test_header_has_shading(self):
        doc = Document()
        params = [ApiParam(name="x")]
        table = build_api_table(doc, params, style)
        for c in range(6):
            assert _get_cell_shading(table.cell(0, c)) == style.HEADER_BG_HEX


# ── Fixed Table Layout ────────────────────────────────────────────────


class TestFixedTableLayout:
    def test_table_width_is_fixed_dxa(self):
        """Table width must be set to exact twips, not auto."""
        doc = Document()
        table = build_generic_info_table(doc, "L", "V", "D", style)
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        tblW = tblPr.find(qn("w:tblW"))
        assert tblW.get(qn("w:type")) == "dxa"
        expected_twips = sum(int(pt * 20) for pt in style.INFO_COLS_PT)
        assert tblW.get(qn("w:w")) == str(expected_twips)

    def test_table_layout_is_fixed(self):
        """Autofit must be disabled via tblLayout type=fixed."""
        doc = Document()
        table = build_generic_info_table(doc, "L", "V", "D", style)
        tbl = table._tbl
        tblPr = tbl.find(qn("w:tblPr"))
        tblLayout = tblPr.find(qn("w:tblLayout"))
        assert tblLayout is not None
        assert tblLayout.get(qn("w:type")) == "fixed"

    def test_grid_columns_match_cell_widths(self):
        """tblGrid gridCol values must match tcW for rendering consistency."""
        doc = Document()
        table = build_generic_info_table(doc, "L", "V", "D", style)
        tbl = table._tbl
        grid = tbl.find(qn("w:tblGrid"))
        grid_widths = [gc.get(qn("w:w")) for gc in grid.findall(qn("w:gridCol"))]
        expected = [str(int(pt * 20)) for pt in style.INFO_COLS_PT]
        assert grid_widths == expected

    def test_no_spacer_paragraphs_after_table(self):
        """Tables must not inject empty spacer paragraphs into the document."""
        doc = Document()
        build_generic_info_table(doc, "L", "V", "D", style)
        # No paragraphs should exist in the document body after a table
        for p in doc.paragraphs:
            assert p.text != "" or (p.style is not None and p.style.name != "Normal"), (
                "Found an empty Normal paragraph — likely a spacer paragraph"
            )
