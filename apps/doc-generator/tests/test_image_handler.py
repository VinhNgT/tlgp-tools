"""Tests for image_handler — embedding and missing image handling."""

from doc_generator.image_handler import insert_image
from docx import Document
from PIL import Image


class TestInsertImage:
    def test_existing_image_embedded(self, tmp_path):
        img = Image.new("RGB", (200, 100), color="blue")
        img_path = tmp_path / "test.png"
        img.save(str(img_path))

        doc = Document()
        insert_image(doc, img_path, full_width=True)
        assert len(doc.inline_shapes) == 1

    def test_missing_image_shows_placeholder(self, tmp_path):
        doc = Document()
        fake_path = tmp_path / "nonexistent.png"
        insert_image(doc, fake_path)

        text = "\n".join(p.text for p in doc.paragraphs)
        assert "Image not found" in text
        assert "nonexistent.png" in text

    def test_natural_size(self, tmp_path):
        img = Image.new("RGB", (200, 100), color="green")
        img_path = tmp_path / "natural.png"
        img.save(str(img_path))

        doc = Document()
        insert_image(doc, img_path, full_width=False)
        assert len(doc.inline_shapes) == 1
